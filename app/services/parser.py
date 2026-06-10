import io

from app.core.exceptions import ParserError, ScannedPDFError, UnsupportedFileTypeError
from app.core.logger import get_logger

logger = get_logger(__name__)


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract plain text from a PDF file using PyMuPDF.

    Raises ParserError if the file is corrupted or password-protected.
    Raises ScannedPDFError if the PDF contains no readable text.
    """
    import fitz  # PyMuPDF

    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
    except Exception as e:
        logger.error(f"Failed to open PDF: {e}")
        raise ParserError("Could not read the file — it may be corrupted")

    if doc.is_encrypted:
        doc.close()
        raise ParserError("Password-protected PDFs are not supported")

    pages_text = []
    for page in doc:
        pages_text.append(page.get_text())
    doc.close()

    text = "\n".join(pages_text).strip()

    if len(text) < 50:
        raise ScannedPDFError("Unable to process the uploaded file. Please try again.")

    logger.info(f"PDF extracted — {len(text)} characters")
    return text


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract plain text from a DOCX file using python-docx.

    Extracts both paragraph text and table content to handle resumes where skills or experience are laid out in tables.
    Raises ParserError if the file is corrupted, empty, or in the old .doc format.
    """
    from docx import Document
    from docx.opc.exceptions import PackageNotFoundError

    try:
        doc = Document(io.BytesIO(file_bytes))
    except PackageNotFoundError:
        raise ParserError(
            "Could not read the file — it may be corrupted or saved in the old .doc format (not .docx)"
        )
    except Exception as e:
        logger.error(f"Failed to open DOCX: {e}")
        raise ParserError("Could not read the file — it may be corrupted")

    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    # extract tables content
    table_cells = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text and cell_text not in table_cells:
                    table_cells.append(cell_text)

    text = "\n".join(paragraphs + table_cells).strip()

    if not text:
        raise ParserError("The uploaded file appears to be empty")

    logger.info(f"DOCX extracted — {len(text)} characters")
    return text


def extract_text(file_bytes: bytes, filename: str) -> str:
    """Route the file to the correct extractor based on its extension.

    Raises UnsupportedFileTypeError if the file is not a PDF or DOCX.
    """
    if not filename or "." not in filename:
        raise UnsupportedFileTypeError("File has no recognizable extension")

    ext = filename.lower().split(".")[-1]

    logger.info(f"Extracting text from {filename} ({len(file_bytes) / 1024:.1f} KB)")

    if ext == "pdf":
        return extract_text_from_pdf(file_bytes)
    elif ext == "docx":
        return extract_text_from_docx(file_bytes)
    else:
        raise UnsupportedFileTypeError(f"Unsupported file type: .{ext}")
